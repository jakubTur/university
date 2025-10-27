import logging
import sys
import os
import fileinput
import matplotlib.pyplot as plt
import numpy as np
import requests
from PIL import Image
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def bookInfo(filename):
    chapter=False
    if not os.path.exists(filename):
        logging.error(f"log file {filename} does not exist")
        sys.exit()
    title, name, content="","",""
    for line in fileinput.input(files=filename,openhook=fileinput.hook_encoded("utf-8")):
        line=line.strip()
        if(title=="" and line.startswith("Title: ")):
            title=line[len("Title: "):]
        if(name=="" and line.startswith("Author: ")):
            name=line[len("Author: "):]
        if(line.startswith("CHAPTER I.")):
            chapter=True
        if(chapter):
            if(line.startswith("CHAPTER II.")):
                break
            else:
                content+=("\n"+line)
    return title, name, content
        
def plotChapter(chapter):
    counterList=[0,0,0,0,0,0,0] #1-10, 11-20, 21-30, 31-40, 41-50, 51-60, 60+
    wordCount=0
    for line in chapter.splitlines():
        if(line.strip()!="" and not line.strip().startswith("*")):
            wordCount+=len(line.split())
        else:
            if(wordCount==0):
                continue
            counted=False
            for i in range(len(counterList)-1):
                if(10*(i+1)>=wordCount):
                    counterList[i]+=1
                    counted=True
                    break
            if(not counted):
                counterList[len(counterList)-1]+=1
            wordCount=0
    counterList[0]-=2 #chapter number and title
    print(counterList)
    threshholds=['1-10', '11-20', '21-30', '31-40', '41-50', '51-60', '60+']
    plt.bar(threshholds, counterList)
    plt.title("Amount of words per paragraph")
    plt.xlabel("words in a paragraph")
    plt.ylabel("amount of paragraphs")
    plt.show()


url_image1 = "https://th.bing.com/th/id/OIP.SBBaTR2Is_gaVs9KUcHLZwHaFJ?rs=1&pid=ImgDetMain"
url_image2 = "https://th.bing.com/th/id/R.fdcb07bf376d27f41aafef2c27b4a98e?rik=z0fh4mU3iTtqRQ&pid=ImgRaw&r=0"
filename_image1 = "image1.jpg"
filename_image2 = "image2.jpg"
final_image_path = "combined_image.jpg"


def download_image(url, filename):
    response = requests.get(url)
    if response.status_code == 200:
        with open(filename, "wb") as f:
            f.write(response.content)
        print(f"Downloaded {filename}")
    else:
        print(f"Failed to download {filename}")


def modify_image1(filename):
    image = Image.open(filename)
    width, height = image.size
    crop_box = (50, 50, width - 50, height - 50)
    cropped = image.crop(crop_box)
    resized = cropped.resize((500, 300))
    return resized


def modify_image2(filename):
    image = Image.open(filename).convert("RGBA")
    rotated = image.rotate(45, expand=True)
    return rotated

def combine_images(base_img, overlay_img):
    base_img = base_img.convert("RGBA")
    overlay_img = overlay_img.convert("RGBA")
    overlay_img = overlay_img.resize((100, 100))
    base_img.paste(overlay_img, (50, 30), overlay_img)
    final_img = base_img.convert("RGB")
    final_img.save("combined_image.jpg")
    print("Combined image saved")

def create_word_doc():
    doc = Document()


    doc.add_heading("Book Report: The Sea Girl", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_heading("by Ray Cummings", level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_heading("Report by: Jakub Turkowski, Dominik Zamuskevic", level=3).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_picture(final_image_path, width=Inches(5))
    doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()
    doc.save("Book.docx")
    print("Word document created: Book.docx")



def run():
    print("This is the lab assignment solution")
    filename="book.txt"
    if(len(sys.argv)>1):
        filename=sys.argv[1]
    title, author, chapter1=bookInfo(filename)
    print(title+"\n"+author+"\n"+chapter1)
    plotChapter(chapter1)

if __name__ == "__main__":
    run()
else:
    logging.error("This file must be run directly, not imported.")
    sys.exit(1)
